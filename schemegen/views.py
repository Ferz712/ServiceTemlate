from django.http import HttpResponse, Http404
from django.shortcuts import render
from django.conf import settings

from rest_framework import viewsets
from rest_framework.response import Response
from rest_framework.decorators import api_view

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docxtpl import DocxTemplate
from datetime import datetime
import zipfile as zp
import json
import re
import os

from .models import Template, TextAlias

def index(request):
    # Да, это очень плохое решение, 
    # Это надо положить в модели
    d = {
    'c0_v1': 'Мировой судья',
    'c0_v2': 'Районный суд',
    'c0_v3': 'Гарнизонный военный суд',
    'c0_v4': 'Областной и приравненный к нему суд',
    'c2_v1': 'Нет',
    'c2_v2': 'Есть',
    'c4_v1': 'Нет',
    'c4_v2': 'Есть',
    'c6_v1': 'Отсутствует право на освобождение от уплаты государственной пошлины',
    'c6_op_v1': 'Исковое заявление имущественного характера (пп.1 п.1 ст.333.19 НК РФ)',
    'c6_op_v2': 'Административное исковое заявление имущественного характера (пп.1 п.1 ст.333.19 НК РФ)',
    'c6_op_v3': 'Исковое заявление имущественного характера, не подлежащее оценке (пп.3 п.1 ст.333.19 НК РФ)',
    'c6_op_v4': 'Исковое заявление неимущественного характера (пп.3 п.1 ст.333.19 НК РФ)',
    'c6_op_v5': 'Исковое заявление о расторжении брака (пп.5 п.1 ст.333.19 НК РФ)',
    'c6_op_v6': 'Административное исковое заявление об оспаривании (полностью или частично) нормативных правовых актов (нормативных актов) государственных органов, Центрального банка РФ, государственных внебюджетных фондов, органов местного самоуправления, государственных корпораций, должностных лиц, а также административного искового заявления об оспаривании ненормативных правовых актов Президента РФ, Совета Федерации Федерального Собрания РФ, Государственной Думы Федерального Собрания РФ, Правительства РФ, Правительственной комиссии по контролю за осуществлением иностранных инвестиций в РФ (пп.6 п.1 ст.333.19 НК РФ)',
    'c6_op_v7': 'Административное исковое заявление об оспаривании актов федеральных органов исполнительной власти, иных федеральных государственных органов, Центрального банка РФ, государственных внебюджетных фондов, содержащих разъяснения законодательства и обладающих нормативными свойствами (пп.6.1 п.1 ст.333.19 НК РФ)',
    'c6_op_v8': 'Административное исковое заявление о признании ненормативного правового акта недействительным и о признании решений и действий (бездействия) государственных органов, органов местного самоуправления, иных органов, должностных лиц незаконными (пп.7 п.1 ст.333.19 НК РФ)',
    'c6_op_v9': 'Исковое заявление о взыскании алиментов (пп.14 п.1 ст.333.19 НК РФ)',
    'c6_op_v10': 'Административное исковое заявление о присуждении компенсации за нарушение права на судопроизводство в разумный срок или права на исполнение судебного акта в разумный срок (пп.15 п.1 ст.333.19 НК РФ)',
    'c6_op_v11': 'Административное исковое заявление о присуждении компенсации за нарушение условий содержания под стражей, содержания в исправительном учреждении (пп.16 п.1 ст.333.19 НК РФ)',
    'c6_v2': 'Есть право на освобождение от уплаты государственной пошлины',
    'c6_v3': 'Категория плательщика',
    'c6_v5': 'Прокуроры - по заявлениям в защиту прав, свобод и законных интересов граждан, неопределенного круга лиц или интересов РФ, субъектов РФ и муниципальных образований (пп. 9 п.1 ст. 333.36 НК РФ)',
    'c6_v6': 'Реабилитированные лица и лица, признанные пострадавшими от политических репрессий, - при обращении по вопросам, возникающим в связи с применением законодательства о реабилитации жертв политических репрессий, за исключением споров между этими лицами и их наследниками (пп. 11 п.1 ст. 333.36 НК РФ)',
    'c6_v7': 'Вынужденные переселенцы и беженцы - при подаче административных исковых заявлений об оспаривании отказа в регистрации ходатайства о признании их вынужденными переселенцами или беженцами (пп. 12 п.1 ст. 333.36 НК РФ)',
    'c6_v8': 'Уполномоченный федеральный орган исполнительной власти по контролю (надзору) в области защиты прав потребителей (его территориальные органы), а также иные федеральные органы исполнительной власти, осуществляющие функции по контролю и надзору в области защиты прав потребителей и безопасности товаров (работ, услуг) (их территориальные органы), органы местного самоуправления, общественные объединения потребителей (их ассоциации, союзы) - по искам, предъявляемым в интересах потребителя, группы потребителей, неопределенного круга потребителей (пп. 13 п.1 ст. 333.36 НК РФ)',
    'c6_v9': 'Уполномоченный по правам человека в РФ - при совершении действий, предусмотренных пп. 1 и 3 п. 1 ст. 29 ФКЗ от 26 февраля 1997 года N 1-ФКЗ "Об Уполномоченном по правам человека в РФ" (пп. 16 п.1 ст. 333.36 НК РФ)',
    'c6_v10': 'Уполномоченный по правам человека в РФ - при совершении действий, предусмотренных пп. 1 и 3 п. 1 ст. 29 ФКЗ от 26 февраля 1997 года N 1-ФКЗ "Об Уполномоченном по правам человека в РФ" (пп. 16 п.1 ст. 333.36 НК РФ)',
    'c6_v11': 'Государственные органы, органы местного самоуправления, выступающие по делам, рассматриваемым Верховным Судом РФ, судами общей юрисдикции, мировыми судьями, в качестве истцов (административных истцов) или ответчиков (административных ответчиков) (пп. 19 п.1 ст. 333.36 НК РФ)',
    'c6_v12': 'Авторы результата интеллектуальной деятельности - по искам о предоставлении им права использования результата интеллектуальной деятельности, исключительное право на который принадлежит другому лицу (принудительная лицензия) (пп. 21 п.1 ст. 333.36 НК РФ)',
    'c6_v13': 'Общественные организации инвалидов, выступающие в качестве истцов (административных истцов) или ответчиков (административных ответчиков) (пп. 1 п.2 ст. 333.36 НК РФ)',
    'c6_v14': 'Истцы (административные истцы) - инвалиды I или II группы (пп. 2 п.2 ст. 333.36 НК РФ)',
    'c6_v15': 'Ветераны боевых действий, ветераны военной службы, обращающиеся за защитой своих прав, установленных законодательством о ветеранах (пп. 3 п.2 ст. 333.36 НК РФ)',
    'c6_v16': 'Истцы - пенсионеры, получающие пенсии, назначаемые в порядке, установленном пенсионным законодательством РФ, - по искам имущественного характера, по административным искам имущественного характера к Пенсионному фонду РФ, негосударственным пенсионным фондам либо к федеральным органам исполнительной власти, осуществляющим пенсионное обеспечение лиц, проходивших военную службу (пп. 5 п.2 ст. 333.36 НК РФ)',
    'c6_v17': 'Физические лица - Герои СССР, Герои РФ и полные кавалеры ордена Славы',
    'c6_v4': 'Категория дела или требования',
    'c6_v18': 'По искам, вытекающим из трудовых правоотношений (пп. 1 п.1 ст. 333.36 НК РФ)',
    'c6_v19': 'По искам о взыскании пособий (пп. 1 п.1 ст. 333.36 НК РФ)',
    'c6_v20': 'По искам о взыскании алиментов (пп. 2 п.1 ст. 333.36 НК РФ)',
    'c6_v21': 'По искам о возмещении вреда, причиненного увечьем или иным повреждением здоровья, а также смертью кормильца (пп. 3 п.1 ст. 333.36 НК РФ)',
    'c6_v22': 'По искам о возмещении имущественного и (или) морального вреда, причиненного преступлением (пп. 4 п.1 ст. 333.36 НК РФ)',
    'c6_v23': 'По искам о возмещении имущественного и (или) морального вреда, причиненного в результате уголовного преследования (пп. 10 п.1 ст. 333.36 НК РФ)',
    'c6_v24': 'При подаче административных исковых заявлений (пп. 7 п.1 ст. 333.36 НК РФ)',
    'c6_v25': 'Истцы - при рассмотрении дел о защите прав и законных интересов ребенка (пп. 15 п.1 ст. 333.36 НК РФ)',
    'c6_v26': 'Истцы - по искам неимущественного характера, связанным с защитой прав и законных интересов инвалидов (пп. 17 п.1 ст. 333.36 НК РФ)',
    'c6_v27': 'Административные истцы - по административным делам о госпитализации гражданина в медицинскую организацию, оказывающую психиатрическую помощь в стационарных условиях, в недобровольном порядке и (или) о психиатрическом освидетельствовании гражданина в недобровольном порядке (пп. 18 п.1 ст. 333.36 НК РФ)',
    'c6_v28': 'Истцы - по искам, связанным с нарушением прав потребителей (пп. 4 п.2 ст. 333.36 НК РФ)',
    'c6_v29': 'Право на освобождение от уплаты государственной пошлины согласно пп. 19 п. 1 ст. 333.36 НК РФ',
    'c7_v1': 'О взыскании задолженности по гражданско-правовому договору',
    'c7_v2': 'О возврате неосновательного обогащения',
    'c7_v3': 'О возврате неосновательно переданного права',
    'c7_v4': 'Да',
    'c7_v5': 'Нет',
    'c7_v6': 'О возврате переданного имущества в натуре',
    'c7_v7': 'Да',
    'c7_v8': 'Нет',
    'c7_v9': 'Да',
    'c7_v10': 'Нет',
    'c7_v11': 'О взыскании стоимости имущества либо по денежным требованиям из неосновательного обогащения',
    'c7_v12': 'Да',
    'c7_v13': 'Нет',
    'c7_v14': 'О возмещении вреда (по обязательствам из причинения вреда, кроме компенсации морального вреда)',
    'c7_v15': 'О возмещении убытков, вытекающих из договоров',
    'c7_v16': 'О признании сделки недействительной',
    'c7_v17': 'О признании права собственности на недвижимое имущество',
    'c7_v18': 'О признании права собственности на движимую вещь',
    'c7_v19': 'О восстановлении в праве (кроме неосновательной передачи права)',
    'c7_v20': 'О восстановлении на работе',
    'c7_v21': 'О восстановлении пропущенного срока',
    'c7_v22': 'Об обязании нотариуса совершить нотариальное действие',
    'c7_v23': 'Об обязании исполнить обязательство в натуре по договору',
    'c7_v24': 'Об обязании устранить недостатки',
    'c7_v25': 'Об обязании передать индивидуально-определенную вещь в соответствии с условиями обязательства (кроме неосновательного обогащения)',
    'c8_v1': 'О взыскании процентов за пользование чужими денежными средствами',
    'c8_v2': 'О взыскании неустойки',
    'c8_v3': 'О взыскании компенсации морального вреда',
    'c8_v4': 'О взыскании расходов на услуги представителя',
    'c8_v5': 'О взыскании иных судебных расходов',
    'c5_v1': 'Указывается',
    'c5_v2': 'Не указывается',
    'c9_v1': 'Да',
    'c9_v2': 'Нет',
    'c10_v1': 'Да',
    'c10_v2': 'Нет',
    'c11_v1': 'Да (Соблюден ли претензионный порядок?)',
    'c11_v2': 'Да',
    'c11_v3': 'Нет',
    'c11_v4': 'Нет'}

    if not request.session or not request.session.session_key:
        request.session.save()
    context = {'full_path':  request.build_absolute_uri()}
    context.update(d)
    return render(request, 'schemegen/index.html', context)

def convertion(ans, session_key):
    print(ans)
    document = Document()
    fname = f'./claims/demo_{session_key}.docx'
    document.save(fname)  
    p = document.add_paragraph()
    p.add_run('{{court}}').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = document.add_paragraph()
    p.add_run("Адрес:_____________________").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT 
    for j in ans[1]:
        res = re.split('\n',j)
        p = document.add_paragraph()
        p.add_run('Истец:').bold = True
        for i in res:
            p.add_run(i)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p = document.add_paragraph()
    if ans[2][0] != '':
        k = 0
        for i in ans[2][1:]:
            res = re.split(r'[1]\d*.|\n[2,3]\d*.|\r\n[2,3]\d*.',i)
            res.pop(0)
            part1 = re.split('\n|\r\n',res[0])
            p.add_run('Представитель Истца:').bold = True
            ans[2][k]=res[1]
            for j in part1:
                p.add_run(j)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p = document.add_paragraph() 
            k+=1
    for i in ans[3]:
        res = re.split('\n',i)
        p = document.add_paragraph()
        p.add_run('Ответчик:').bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for j in res:            
            p.add_run(j)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT      
            p = document.add_paragraph()
    if ans[4][0] != '':
        for i in ans[4][1:]:
            p.add_run('Третье лицо: ').bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            res = re.split('\n',i)
            for j in res:
                p.add_run(j)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT 
                p = document.add_paragraph()
    if ans[5] != ['']:
        p = document.add_paragraph()
        p.add_run('{{price_isk}}').bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if ans[6] != ['оплачивается']:
        p = document.add_paragraph()
        p.add_run('Государственная пошлина: ').bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = document.add_paragraph()
        p.add_run('{{poshlina}}')
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p = document.add_paragraph()
        p.add_run('Государственная пошлина: ').bold = True
        p.add_run('_________')
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = document.add_paragraph()
    p.add_run("Исковое заявление").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    k = 0
    l = []
    for i in ans[7]:
        res = re.split(r'[1]\d*.|\n[2,3]\d*.|\r\n[2,3]\d*.',i)
        res.pop(0)
        l.append(res)
        if i == ans[7][-1]:
            ans[7] = l
            break
    if ans[8]!=['']:
        l=[]
        for i in ans[8]:
            res = re.split(r'[1]\d*.|\n[2,3]\d*.',i)
            res.pop(0)
            l.append(res)
            if i == ans[8][-1]:
                ans[8] = l
                break
    p = document.add_paragraph()
    s = ''
    for i in ans[7]:
        s += i[0]+', '
    s = s[:-2]
    p.add_run('о ' + s).italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph("ПРОШУ:")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    k = 0
    for i in ans[7]:
        res = re.split(r'\n|\r\n', i[1])
        for j in res:
            k+=1
            document.add_paragraph('%i.' %k + j)
    if ans[8]!= ['']:
        l = []
        for i in ans[8]:
            l.append(i[0])
        res = set(l)
        for i in res:
            k+=1
            document.add_paragraph('%i.' %k + i)
    if ans[11]!= ['']:
        k+=1
        document.add_paragraph('%i.' %k + '{{potrebiteli}}')
    p = document.add_paragraph()
    p.add_run('Приложение:').bold = True
    if ans[6] == ['оплачивается']:
        document.add_paragraph('Платежное поручение №___ от «__»______ ____ г., подтверждающее уплату государственной пошлины.', style = 'List Number')
    document.add_paragraph('Копия уведомления о вручении или иные документы, подтверждающие направление другим лицам, участвующим в деле, копий искового заявления и приложенных к нему документов, которые у других лиц, участвующих в деле, отсутствуют;', style = 'List Number')
    document.add_paragraph('Иные документы, на которых Истец обосновывает свои требования;', style = 'List Number')
    if ans[2][0]!= '':
        document.add_paragraph('{{complainant}}', style = 'List Number')
    for i in ans[7]:
        if i[2] != '':
            res = re.split('\n|\r\n',i[2])
            for j in res:
                document.add_paragraph(j, style = 'List Number')
    if ans[8] != ['']:
        l = []
        for i in ans[8]:
            l.append(i[1])
        res = set(l)
        for i in res:
            document.add_paragraph(i, style = 'List Number')
    if ans[9] != ['']:
        document.add_paragraph('{{regulation}}', style = 'List Number')
    if ans[10] != ['']:
        document.add_paragraph('{{peace}}', style = 'List Number')
    document.add_paragraph('« » ________ _____г. \t\t\t\t\t\t_____________ (________________)')
    document.save(fname)
    document = DocxTemplate(fname)
    context = {'court':ans[0][0], 'complainant':ans[2][0], 'price_isk':ans[5][0],'poshlina':ans[6][0], 'regulation':ans[9][0], 'peace':ans[10][0], 'potrebiteli':ans[11][0]}
    document.render(context)
    document.save(fname)

def get_text(request):
    req = dict(request.POST)
    # print(req)

    res = {}
    for choice in req:
        if len(choice.split('-')) != 2:
            continue

        i = int(choice.split('-')[1])
        answer = []

        if choice == 'choice-7' and 'c7-v2' in req[choice]:
            op1 = req.get('choice-7-op1')
            op2 = req.get('choice-7-op2')
            op3 = req.get('choice-7-op3')
            op4 = req.get('choice-7-op4')

            if op1 is not None:
                answer.append(TextAlias.objects.get(html_id=op1[0]).text)
            if op2 is not None and op3 is not None:
                ans0 = TextAlias.objects.get(html_id=op2[0]).text
                ans1 = TextAlias.objects.get(html_id=op3[0]).text
                answer.append(ans0 + '\n' + ans1)
            if op4 is not None:
                answer.append(TextAlias.objects.get(html_id=op4[0]).text)

            req['choice-7'].remove('c7-v2')
        
        for j in req[choice]:
            answer.append(TextAlias.objects.get(html_id=j).text)
        
        res[i] = answer
    
    convertion(res, request.session.session_key)

    return HttpResponse("Success")

def download(request):
    session_key = request.session.session_key
    os.system(f'doc2pdf --output=./claims/template_{session_key} ./claims/demo_{session_key}.docx')

    path = f'./claims/template_{session_key}.pdf'
    file_path = os.path.join(settings.MEDIA_ROOT, path)
    # print(file_path)
    if os.path.exists(file_path):
        with open(file_path, 'rb') as fh:
            response = HttpResponse(fh.read(), content_type="application/pdf")
            response['Content-Disposition'] = 'inline; filename=' + os.path.basename(file_path)
            return response
    raise Http404

def download_doc(request):
    # os.system('doc2pdf --output=template demo.docx')
    session_key = request.session.session_key
    path = f'./claims/demo_{session_key}.docx'
    file_path = os.path.join(settings.MEDIA_ROOT, path)
    dt = datetime.now()
    # print(file_path)
    if os.path.exists(file_path):
        with open(file_path, 'rb') as fh:
            response = HttpResponse(fh.read(), content_type="application/msdocx")
            response['Content-Disposition'] = 'inline; filename=' + f'Generated_claim_{dt.strftime("%Y%m%d%H%M%S")}.docx'
            return response
    raise Http404
